"""FOOUND Synthesis Runner v1.

Contract: Synthesis Runner v1 plan + Amendment A + Amendment A.1 + Migration 008.

The runner is deliberately THIN. It discovers, reads, interprets, validates,
and calls the doors. It never reimplements database truth:

  claim      -> claim_synthesis_batch(job)        (the only way in)
  settle     -> settle_synthesis_results(job,...) (the only successful ending:
                item verdicts + memory + sufficiency + lifecycle, atomically)
  abort      -> finalize_synthesis(job,'failed')  (only when no validated
                per-item verdicts exist: over-budget batch, storage/model
                death, engine error, janitor)

The database owns claim, settlement, withdrawal-wins, provenance acceptance,
the sufficiency outcome, and lifecycle completion. The sufficiency POLICY
constants live here (reviewed code); the computation lives in the door.

Privacy: the engine repo is PUBLIC and GitHub Actions logs are public.
Logging is ids / counts / enums / timings ONLY. Never evidence text, labels,
storage paths, memory statements, prompts, or model output. The single log()
helper is the only logging path.

No model-authored client copy: everything the client can ever see is either
baked into migrations 007/008 or listed in FROZEN_CLIENT_COPY below.
"""

from __future__ import annotations

import json
import logging
import re
import time
from dataclasses import dataclass, field
from typing import Optional

# ---------------------------------------------------------------------------
# Reviewed constants (pilot calibration — provisional, changed only by commit)
# ---------------------------------------------------------------------------

MODEL_ID = "claude-sonnet-4-5-20250929"  # full immutable snapshot id, never an alias
MODEL_MAX_OUTPUT_TOKENS = 4096
MODEL_MALFORMED_RETRIES = 1     # one retry with the validation error appended
MODEL_API_RETRIES = 2           # infrastructure retries with backoff
STORAGE_FETCH_RETRIES = 2

PER_ITEM_CHAR_CAP = 150_000     # above: honest per-item failure (too_large)
BATCH_INPUT_CHAR_BUDGET = 400_000  # above: honest whole-job abort
MAX_STATEMENTS_PER_JOB = 40
RECONCILIATION_ROW_CAP = 200    # comparison context (active+tension) MAY be capped
RETRACTED_FETCH_LIMIT = 1000    # suppression context must be COMPLETE: if more
                                # retracted rows exist than this, synthesis
                                # FAILS CLOSED rather than run while knowingly
                                # missing client retractions

JANITOR_STALE_MINUTES = 30

# Sufficiency policy: VALUES here, COMPUTATION in settle_synthesis_results().
POLICY = {
    "min_grounded": 5,
    "require_record": True,
    "require_self": True,
    "require_direction": True,
    "max_failed_ratio": 0.5,
}

ALLOWED_MIME = {
    "text/plain": "text",
    "text/markdown": "text",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}

# The ONLY client-visible strings this module may ever pass to a door.
# Item-level copy lives in the 008 taxonomy (codes below); job-level default
# copy lives in 007's finalize. This dict is frozen; the model has no path
# to it or to any door error argument.
FROZEN_CLIENT_COPY = {
    "batch_too_large": (
        "FOOUND can't read this much in one pass. "
        "Remove these items, then add a smaller set and try again."
    ),
}

FAILURE_CODES = {"unreadable", "no_text_pdf", "too_large"}  # 008 closed taxonomy

UUID_RE = re.compile(r"^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$")

log = logging.getLogger("synthesis_runner")


# ---------------------------------------------------------------------------
# Door / transport abstraction. Two implementations: REST (production) and
# direct Postgres (local harness). Both call the SAME database functions —
# nothing here interprets contract semantics.
# ---------------------------------------------------------------------------

class DoorError(Exception):
    """A named refusal raised by a database door (e.g. 'job_not_queued')."""

    def __init__(self, name: str):
        super().__init__(name)
        self.name = name


class Db:
    """Transport interface to the database. All methods are thin."""

    def oldest_queued_synthesize_job(self) -> Optional[dict]:
        raise NotImplementedError

    def stale_running_synthesize_jobs(self, stale_minutes: int) -> list[dict]:
        raise NotImplementedError

    def claim(self, job_id: str) -> dict:
        raise NotImplementedError

    def settle(self, job_id: str, results: dict, policy: dict) -> dict:
        raise NotImplementedError

    def finalize_failed(self, job_id: str, error: Optional[str]) -> dict:
        raise NotImplementedError

    def evidence_rows(self, item_ids: list[str]) -> list[dict]:
        raise NotImplementedError

    def comparison_memory(self, agent_id: str) -> list[dict]:
        """Active + tension rows for the comparison context — deliberately
        capped at RECONCILIATION_ROW_CAP. Read-only; never citable."""
        raise NotImplementedError

    def retracted_memory(self, agent_id: str) -> list[dict]:
        """Retracted rows for the suppression context. Fetched SEPARATELY
        from the comparison context so the cap there can never silently
        truncate retractions. Fetches up to RETRACTED_FETCH_LIMIT + 1 rows;
        the runner FAILS CLOSED if the sentinel row appears."""
        raise NotImplementedError


class RestDb(Db):
    """Production transport: Supabase PostgREST with the service key."""

    def __init__(self, base_url: str, service_key: str):
        import requests  # local import: harness never needs it

        self._requests = requests
        self.base = base_url.rstrip("/")
        self.headers = {
            "apikey": service_key,
            "Authorization": f"Bearer {service_key}",
            "Content-Type": "application/json",
        }

    def _rpc(self, fn: str, payload: dict):
        r = self._requests.post(
            f"{self.base}/rest/v1/rpc/{fn}", headers=self.headers,
            data=json.dumps(payload), timeout=60,
        )
        if r.status_code >= 400:
            try:
                msg = r.json().get("message", "")
            except Exception:
                msg = ""
            raise DoorError(msg or f"http_{r.status_code}")
        return r.json()

    def _select(self, path: str) -> list[dict]:
        r = self._requests.get(
            f"{self.base}/rest/v1/{path}", headers=self.headers, timeout=60,
        )
        r.raise_for_status()
        return r.json()

    def oldest_queued_synthesize_job(self):
        rows = self._select(
            "jobs?type=eq.synthesize&status=eq.queued"
            "&select=id,agent_id,requested_at&order=requested_at.asc&limit=1"
        )
        return rows[0] if rows else None

    def stale_running_synthesize_jobs(self, stale_minutes):
        import datetime as _dt

        # URL-safe UTC format: a '+00:00' offset would put a raw '+' in the
        # query string, which HTTP decodes as a space -> PostgREST 400.
        # (Found live in Fire #1 run 1; W6 now forbids unsafe characters.)
        cutoff = (
            _dt.datetime.now(_dt.timezone.utc) - _dt.timedelta(minutes=stale_minutes)
        ).strftime("%Y-%m-%dT%H:%M:%SZ")
        return self._select(
            "jobs?type=eq.synthesize&status=eq.running"
            f"&started_at=lt.{cutoff}&select=id,agent_id,started_at"
        )

    def claim(self, job_id):
        return self._rpc("claim_synthesis_batch", {"p_job": job_id})

    def settle(self, job_id, results, policy):
        return self._rpc(
            "settle_synthesis_results",
            {"p_job": job_id, "p_results": results, "p_policy": policy},
        )

    def finalize_failed(self, job_id, error):
        return self._rpc(
            "finalize_synthesis",
            {"p_job": job_id, "p_outcome": "failed", "p_error": error},
        )

    def evidence_rows(self, item_ids):
        # bare UUIDs in the in-list: no quotes, no URL-unsafe characters
        ids = ",".join(item_ids)
        return self._select(
            f"evidence_items?id=in.({ids})"
            "&select=id,agent_id,kind,label,storage_path,body,mime_type,"
            "byte_size,status,submitted_in"
        )

    def comparison_memory(self, agent_id):
        return self._select(
            f"memory?agent_id=eq.{agent_id}&status=in.(active,tension)"
            "&select=id,layer,statement,provenance,status"
            f"&order=created_at.asc&limit={RECONCILIATION_ROW_CAP}"
        )

    def retracted_memory(self, agent_id):
        # Separate query: the comparison cap can never truncate retractions.
        # limit is FETCH_LIMIT + 1 — the extra row is the fail-closed sentinel.
        return self._select(
            f"memory?agent_id=eq.{agent_id}&status=eq.retracted"
            "&select=id,statement,status"
            f"&order=created_at.asc&limit={RETRACTED_FETCH_LIMIT + 1}"
        )


class PgDb(Db):
    """Harness transport: direct Postgres, calling the same doors."""

    def __init__(self, dsn: str):
        import psycopg2
        import psycopg2.extras

        self._psycopg2 = psycopg2
        self._extras = psycopg2.extras
        self.conn = psycopg2.connect(dsn)
        self.conn.autocommit = True

    def _rows(self, sql: str, args=()):
        with self.conn.cursor(cursor_factory=self._extras.RealDictCursor) as cur:
            cur.execute(sql, args)
            return [dict(r) for r in cur.fetchall()]

    def _door(self, sql: str, args=()):
        try:
            with self.conn.cursor() as cur:
                cur.execute(sql, args)
                return cur.fetchone()[0]
        except self._psycopg2.Error as e:
            raise DoorError((e.diag.message_primary or "").strip()) from e

    def oldest_queued_synthesize_job(self):
        rows = self._rows(
            "select id::text, agent_id::text, requested_at from jobs "
            "where type='synthesize' and status='queued' "
            "order by requested_at asc limit 1"
        )
        return rows[0] if rows else None

    def stale_running_synthesize_jobs(self, stale_minutes):
        return self._rows(
            "select id::text, agent_id::text, started_at from jobs "
            "where type='synthesize' and status='running' "
            "and started_at < now() - (%s || ' minutes')::interval",
            (stale_minutes,),
        )

    def claim(self, job_id):
        return self._door("select claim_synthesis_batch(%s)", (job_id,))

    def settle(self, job_id, results, policy):
        return self._door(
            "select settle_synthesis_results(%s, %s::jsonb, %s::jsonb)",
            (job_id, json.dumps(results), json.dumps(policy)),
        )

    def finalize_failed(self, job_id, error):
        return self._door(
            "select finalize_synthesis(%s, 'failed', %s)", (job_id, error)
        )

    def evidence_rows(self, item_ids):
        return self._rows(
            "select id::text, agent_id::text, kind, label, storage_path, body, "
            "mime_type, byte_size, status, submitted_in::text "
            "from evidence_items where id = any(%s::uuid[])",
            (item_ids,),
        )

    def comparison_memory(self, agent_id):
        return self._rows(
            "select id::text, layer, statement, provenance, status from memory "
            "where agent_id=%s and status in ('active','tension') "
            "order by created_at asc limit %s",
            (agent_id, RECONCILIATION_ROW_CAP),
        )

    def retracted_memory(self, agent_id):
        return self._rows(
            "select id::text, statement, status from memory "
            "where agent_id=%s and status = 'retracted' "
            "order by created_at asc limit %s",
            (agent_id, RETRACTED_FETCH_LIMIT + 1),
        )


# ---------------------------------------------------------------------------
# Storage and model abstractions.
# ---------------------------------------------------------------------------

class Storage:
    def fetch(self, path: str) -> bytes:
        raise NotImplementedError


class SupabaseStorage(Storage):
    def __init__(self, base_url: str, service_key: str, bucket: str = "feeds"):
        import requests

        self._requests = requests
        self.base = base_url.rstrip("/")
        self.bucket = bucket
        self.headers = {
            "apikey": service_key,
            "Authorization": f"Bearer {service_key}",
        }

    def fetch(self, path: str) -> bytes:
        last = None
        for attempt in range(STORAGE_FETCH_RETRIES + 1):
            try:
                r = self._requests.get(
                    f"{self.base}/storage/v1/object/{self.bucket}/{path}",
                    headers=self.headers, timeout=120,
                )
                r.raise_for_status()
                return r.content
            except Exception as e:  # noqa: BLE001 — retried, then surfaced
                last = e
                time.sleep(2 * (attempt + 1))
        raise last


class ModelClient:
    def complete(self, system: str, user: str) -> str:
        raise NotImplementedError


class AnthropicModel(ModelClient):
    def __init__(self, api_key: str):
        import anthropic

        self.client = anthropic.Anthropic(api_key=api_key)

    def complete(self, system: str, user: str) -> str:
        last = None
        for attempt in range(MODEL_API_RETRIES + 1):
            try:
                msg = self.client.messages.create(
                    model=MODEL_ID,
                    max_tokens=MODEL_MAX_OUTPUT_TOKENS,
                    system=system,
                    messages=[{"role": "user", "content": user}],
                )
                return "".join(
                    b.text for b in msg.content if getattr(b, "type", "") == "text"
                )
            except Exception as e:  # noqa: BLE001 — retried, then surfaced
                last = e
                time.sleep(5 * (attempt + 1))
        raise last


# ---------------------------------------------------------------------------
# Parsers — the four contract types only. No OCR, no images, no truncation.
# ---------------------------------------------------------------------------

class ParseFailure(Exception):
    def __init__(self, code: str):
        assert code in FAILURE_CODES
        super().__init__(code)
        self.code = code


def parse_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def parse_pdf(data: bytes) -> str:
    import io

    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as e:  # malformed PDF
        raise ParseFailure("unreadable") from e
    if not text.strip():
        raise ParseFailure("no_text_pdf")  # scanned/imagey: we have no OCR
    return text


def parse_docx(data: bytes) -> str:
    import io

    import docx

    try:
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as e:
        raise ParseFailure("unreadable") from e


def extract_content(row: dict, storage: Storage) -> str:
    """Return the full extracted text of one evidence item, or ParseFailure."""
    if row["kind"] == "text":
        content = row["body"] or ""
    else:
        try:
            data = storage.fetch(row["storage_path"])
        except Exception as e:
            raise ParseFailure("unreadable") from e
        if len(data) != row["byte_size"]:
            raise ParseFailure("unreadable")  # stored object contradicts the row
        family = ALLOWED_MIME.get(row["mime_type"])
        if family is None:
            raise ParseFailure("unreadable")
        if family == "text":
            content = parse_text(data)
        elif family == "pdf":
            content = parse_pdf(data)
        else:
            content = parse_docx(data)
    if len(content) > PER_ITEM_CHAR_CAP:
        raise ParseFailure("too_large")  # never truncate: fail honestly
    return content


# ---------------------------------------------------------------------------
# Prompt and strict output validation.
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = """You are the reading engine of FOOUND, a career agent. You convert a client's evidence into grounded memory statements. Output ONLY a JSON object, no prose, no markdown fences.

Rules — absolute:
- Every statement must be grounded in the CURRENT EVIDENCE items given below, cited by their ids. If the evidence does not establish it, it does not exist.
- Never use outside world knowledge about employers, markets, salaries, or people. Never invent, estimate, or repair any date, number, name, or title.
- provenance: "stated" = the evidence says it directly; "extracted" = derivable purely by reading the documents (e.g. a duration from stated dates); "inferred" = a conservative cross-document inference a careful human reader would draw. Nothing else.
- layer: "record" = verifiable career facts; "self" = the person's own account of themselves; "model" = your synthesized understanding. Nothing else.
- You make NO judgment about whether the evidence is sufficient, ready, or good. That is not your role.
- Contradictions must surface, never be averaged, softened, or silently resolved.
- EXISTING MEMORY (if present) is comparison context ONLY. It may never be cited as evidence, and may never cause a new claim that current evidence does not support. Use it only to detect: (a) statements that duplicate existing memory -> report in "reinforcements"; (b) conflicts between current evidence and existing memory -> report in "contradictions" with kind "existing". Only rows marked (active/...) may be referenced by id; rows marked (tension/...) are context only.
- CLIENT-RETRACTED UNDERSTANDING (if present) is suppression context ONLY: the client has struck those beliefs. Never reassert any of them in ANY wording — not as a statement, not as a contradiction, not as a reinforcement — even if the current evidence appears to support them. They are not beliefs and not evidence.

Output JSON shape:
{
 "statements": [{"layer": "...", "statement": "<=1000 chars", "provenance": "...",
                 "evidence": ["<current item id>", ...], "is_direction": bool,
                 "handle": "<optional 1-3 word editorial icon, <=24 chars>"}],
 "contradictions": [{"kind": "batch" | "existing",
                     "a": "<one reading>", "b": "<the conflicting reading>",
                     "evidence": ["<current item ids behind the conflict>"],
                     "existing_memory_id": "<id, only when kind=existing>"}],
 "reinforcements": [{"existing_memory_id": "<id>",
                     "evidence": ["<current item ids that re-support it>"],
                     "is_direction": bool}],
 "unknowns": ["things the evidence approaches but does not establish"]
}
"is_direction" is true only for statements about what the person wants next.
"handle" is optional presentation metadata: an editorial icon in the person's career story, not a compressed summary of the statement. Statement holds precision. Handle holds signal. It is not evidence, not a belief, and never a substitute for the statement.
Judge each handle on four tests: SOUND (rings in the ear), SIGHT (strong in giant type), STORY (adds to the career narrative), TRUTH (supported by the statement; no new claim).
Optimize for: strong proper nouns; powerful career ideas; memorable themes; simple language; one clear subject; visual strength; narrative range across the full set.
Avoid: résumé shorthand; miniature job descriptions; generic UI labels; overly literal compression; clever but weak abstractions; phrases that only make sense after opening the row.
Do not author or judge handles as one isolated label per row. Author and review the entire handle set in this output together as one editorial composition: rhythm, variety, recognizable names, themes, ambition, and point of view. Signal mix across the set: brands -> scale -> craft -> leadership -> worldview -> ambition. Avoid duplicate handles and overloading one signal type.
Structural: a positive career-story Memory may receive a handle; a stated filter or exclusion may omit handle (NULL). NULL does not mean unimportant. Do not invent a positive icon for a filter.
Mechanical: 1-3 words, target <=24 chars, no terminal punctuation, distinct within this output where possible. Omit the field if the handle would be malformed.
Your response MUST begin with the character { and end with the character }."""


def _strip_fences(raw: str) -> str:
    """Deterministically unwrap ONE outer markdown code fence, if present.

    Transport normalization only — models sometimes fence JSON despite
    instructions (observed live: Fire #2 run 3, reason=not_json twice).
    Anything beyond a single clean fence still fails validation honestly."""
    s = raw.strip()
    if s.startswith("```") and s.endswith("```"):
        first_nl = s.find("\n")
        if first_nl != -1:
            s = s[first_nl + 1 : -3].strip()
    return s


def build_user_prompt(
    items: list[dict], contents: dict, memory: list[dict],
    retracted: list[dict] | None = None,
) -> str:
    parts = ["CURRENT EVIDENCE (citable):"]
    for it in items:
        parts.append(
            f'\n--- item id: {it["id"]} · label: {it["label"]} · kind: {it["kind"]} ---'
        )
        parts.append(contents[it["id"]])
    if memory:
        parts.append(
            "\n=== EXISTING MEMORY — COMPARISON CONTEXT ONLY, NEVER CITABLE ==="
        )
        for m in memory:
            parts.append(
                f'[{m["id"]}] ({m["status"]}/{m["layer"]}/{m["provenance"]}) {m["statement"]}'
            )
        parts.append("=== END EXISTING MEMORY ===")
    else:
        parts.append("\n(no existing memory)")
    # 009 fence: retracted beliefs — suppression context ONLY. Deliberately
    # id-less so they are structurally unreferencable: not citable, not
    # reinforceable, not usable in contradictions. They are not beliefs and
    # not evidence; they exist only so the model can avoid repeating them.
    if retracted:
        parts.append(
            "\n=== CLIENT-RETRACTED UNDERSTANDING — DO NOT REASSERT ===\n"
            "These are not evidence and may never support a new claim. They "
            "exist only to prevent the engine from regenerating understanding "
            "the client explicitly rejected. Never reassert any of them, in "
            "any wording, as a statement, contradiction, or reinforcement — "
            "even if the current evidence appears to support one: the "
            "client's retraction stands."
        )
        for m in retracted:
            parts.append(f"- {m['statement']}")
        parts.append("=== END CLIENT-RETRACTED UNDERSTANDING ===")
    return "\n".join(parts)


class ValidationError(Exception):
    pass


def _norm(statement: str) -> str:
    return re.sub(r"\s+", " ", statement).strip().lower()


HANDLE_MAX_CHARS = 40  # defensive ceiling (mirrors the DB CHECK); prompt targets <=24


def _soft_handle(value) -> str | None:
    """010: fail-soft handle normalization. Presentation metadata only —
    a malformed, empty, or oversized handle degrades to None and NEVER
    fails a synthesis. Never participates in _norm/suppression/duplicate
    logic (those remain statement-based by contract)."""
    if not isinstance(value, str):
        return None
    v = value.strip()
    if not v or len(v) > HANDLE_MAX_CHARS:
        return None
    return v


def validate_and_map(
    raw: str, readable_ids: set[str], existing: list[dict],
    retracted: list[dict] | None = None,
) -> tuple[list[dict], list[dict]]:
    """Validate model output strictly; map it to 008 memory/reinforce entries.

    Returns (memory_entries, reinforce_entries). Raises ValidationError with a
    terse machine-readable reason (safe to log AND to feed back on retry —
    it never contains evidence or model text)."""
    try:
        out = json.loads(_strip_fences(raw))
    except Exception:
        raise ValidationError("not_json") from None
    if not isinstance(out, dict):
        raise ValidationError("not_object")
    if set(out.keys()) - {"statements", "contradictions", "reinforcements", "unknowns"}:
        raise ValidationError("unknown_top_key")
    stmts = out.get("statements", [])
    contras = out.get("contradictions", [])
    reinfs = out.get("reinforcements", [])
    unknowns = out.get("unknowns", [])
    if not all(isinstance(x, list) for x in (stmts, contras, reinfs, unknowns)):
        raise ValidationError("bad_shape")
    if len(stmts) > MAX_STATEMENTS_PER_JOB:
        raise ValidationError("too_many_statements")
    # 009: references (reinforcement targets, contradiction existing_memory_id)
    # must resolve to ACTIVE rows only. Tension rows are context, not targets
    # (the door would refuse them); retracted rows never reach this function
    # and carry no ids in the prompt at all.
    active_ids = {m["id"] for m in existing if m["status"] == "active"}
    existing_by_norm = {
        _norm(m["statement"]): m["id"] for m in existing if m["status"] == "active"
    }
    # 009 correction 3: deterministic exact-suppression guard across EVERY
    # model-output channel. Same normalization as the SQL guard — case and
    # whitespace only, never fuzzy. A reassertion is REFUSED (validation
    # error -> documented retry -> honest abort), as a statement or as
    # either side of a contradiction. Reinforcement is already structurally
    # active-id-only; retracted rows carry no ids in the prompt at all.
    retracted_norms = {_norm(m["statement"]) for m in (retracted or [])}

    def check_citations(ev) -> list[str]:
        if not isinstance(ev, list) or not ev:
            raise ValidationError("bad_citations")
        seen = []
        for e in ev:
            if not isinstance(e, str) or not UUID_RE.match(e.lower()):
                raise ValidationError("bad_citation_id")
            el = e.lower()
            if el not in readable_ids:
                raise ValidationError("citation_outside_batch")
            if el not in seen:
                seen.append(el)
        return seen

    memory_entries: list[dict] = []
    reinforce_entries: list[dict] = []
    seen_norms: dict[str, dict] = {}
    reinforced_targets: set[str] = set()

    for s in stmts:
        if not isinstance(s, dict) or set(s.keys()) - {
            "layer", "statement", "provenance", "evidence", "is_direction",
            "handle",  # 010: optional presentation handle
        }:
            raise ValidationError("bad_statement_shape")
        layer = s.get("layer")
        text = s.get("statement")
        prov = s.get("provenance")
        is_dir = s.get("is_direction", False)
        hnd = _soft_handle(s.get("handle"))  # 010: fail-soft, may be None
        if layer not in ("record", "self", "model"):
            # 'behavior' and anything else refused from this producer
            raise ValidationError("layer_not_allowed")
        if prov not in ("stated", "extracted", "inferred"):
            raise ValidationError("provenance_not_allowed")
        if not isinstance(text, str) or not 1 <= len(text) <= 1000:
            raise ValidationError("bad_statement_text")
        if not isinstance(is_dir, bool):
            raise ValidationError("bad_is_direction")
        cites = check_citations(s.get("evidence"))
        norm = _norm(text)
        if norm in retracted_norms:
            raise ValidationError("reasserted_retracted")
        if norm in seen_norms:
            # intra-payload duplicate: merge citations into the first, never
            # send a duplicate to the door (the door would refuse it)
            first = seen_norms[norm]
            first["evidence"] = list(dict.fromkeys(first["evidence"] + cites))
            first["is_direction"] = first.get("is_direction", False) or is_dir
            # 010: handle of the FIRST occurrence wins (deterministic); a
            # duplicate's handle is discarded with its duplicate statement.
            continue
        if norm in existing_by_norm:
            # duplicate of existing active memory: MUST go through reinforce
            target = existing_by_norm[norm]
            if target not in reinforced_targets:
                reinforced_targets.add(target)
                reinforce_entries.append(
                    {"memory": target, "evidence": cites, "is_direction": is_dir}
                )
            continue
        entry = {
            "layer": layer, "statement": text, "provenance": prov,
            "evidence": cites, "is_direction": is_dir,
        }
        if hnd is not None:
            entry["handle"] = hnd  # 010: absent key -> NULL at the door
        seen_norms[norm] = entry
        memory_entries.append(entry)

    for c in contras:
        if not isinstance(c, dict) or set(c.keys()) - {
            "kind", "a", "b", "evidence", "existing_memory_id"
        }:
            raise ValidationError("bad_contradiction_shape")
        kind = c.get("kind")
        if kind not in ("batch", "existing"):
            raise ValidationError("bad_contradiction_kind")
        a, b = c.get("a"), c.get("b")
        if not isinstance(a, str) or not isinstance(b, str) or not a or not b:
            raise ValidationError("bad_contradiction_text")
        if _norm(a) in retracted_norms or _norm(b) in retracted_norms:
            raise ValidationError("contradiction_reasserts_retracted")
        if kind == "existing":
            mid = c.get("existing_memory_id")
            if not isinstance(mid, str) or mid not in active_ids:
                raise ValidationError("bad_existing_memory_id")
        cites = check_citations(c.get("evidence"))
        text = f"Tension: {a} / {b}"[:1000]
        memory_entries.append(
            {
                "layer": "model", "statement": text, "provenance": "extracted",
                "evidence": cites, "tension": True,
            }
        )

    for rf in reinfs:
        if not isinstance(rf, dict) or set(rf.keys()) - {
            "existing_memory_id", "evidence", "is_direction"
        }:
            raise ValidationError("bad_reinforcement_shape")
        mid = rf.get("existing_memory_id")
        if not isinstance(mid, str) or mid not in active_ids:
            raise ValidationError("bad_existing_memory_id")
        is_dir = rf.get("is_direction", False)
        if not isinstance(is_dir, bool):
            raise ValidationError("bad_is_direction")
        cites = check_citations(rf.get("evidence"))
        if mid in reinforced_targets:
            continue  # door refuses duplicate targets; first wins
        reinforced_targets.add(mid)
        reinforce_entries.append(
            {"memory": mid, "evidence": cites, "is_direction": is_dir}
        )

    for u in unknowns:
        if not isinstance(u, str):
            raise ValidationError("bad_unknown")

    return memory_entries, reinforce_entries


# ---------------------------------------------------------------------------
# The runner.
# ---------------------------------------------------------------------------

@dataclass
class RunReport:
    """Machine-readable outcome of one invocation. Counts and enums only."""

    action: str = "none"          # none|empty|refused|race|settled|aborted|error
    job_id: Optional[str] = None
    outcome: Optional[str] = None
    janitor_finalized: int = 0
    detail: dict = field(default_factory=dict)


class Runner:
    def __init__(self, db: Db, storage: Storage, model: ModelClient):
        self.db = db
        self.storage = storage
        self.model = model

    # -- janitor ------------------------------------------------------------
    def janitor(self) -> int:
        finalized = 0
        for job in self.db.stale_running_synthesize_jobs(JANITOR_STALE_MINUTES):
            try:
                fin = self.db.finalize_failed(job["id"], None)  # 007 default copy
                log.info(
                    "janitor finalized job=%s swept=%s",
                    job["id"], fin.get("swept_reading_items"),
                )
                finalized += 1
            except DoorError as e:
                # job_not_running = another actor beat us; anything else is
                # logged by name and skipped (never crashes the run)
                log.info("janitor skip job=%s door=%s", job["id"], e.name)
        return finalized

    # -- one invocation -----------------------------------------------------
    def run_once(self) -> RunReport:
        report = RunReport()
        report.janitor_finalized = self.janitor()

        job = self.db.oldest_queued_synthesize_job()
        if job is None:
            log.info("no queued synthesize jobs")
            return report
        report.job_id = job["id"]
        log.info("discovered job=%s agent=%s", job["id"], job["agent_id"])

        # CLAIM — the only door in.
        try:
            claim = self.db.claim(job["id"])
        except DoorError as e:
            if e.name == "job_not_queued":
                log.info("lost claim race job=%s", job["id"])
                report.action = "race"
                return report
            log.error("claim door error job=%s door=%s", job["id"], e.name)
            report.action = "error"
            report.detail["door"] = e.name
            return report

        status = claim.get("status")
        if status == "refused":
            # claim persisted the terminal failed job itself; nothing to do
            log.info("claim refused job=%s reason=%s", job["id"], claim.get("reason"))
            report.action = "refused"
            return report
        if status == "empty":
            # claim persisted the honest zero-item failure itself
            log.info("claim empty job=%s", job["id"])
            report.action = "empty"
            return report
        if status != "claimed":
            log.error("claim unexpected status job=%s", job["id"])
            report.action = "error"
            return report

        item_ids = [i.lower() for i in claim["items"]]
        log.info("claimed job=%s items=%d", job["id"], len(item_ids))

        try:
            return self._process(job, item_ids, report)
        except Exception:
            # Engine bug or environment death with the job running: attempt the
            # abort door so the client is never stuck; janitor is the backstop.
            log.exception("processing error job=%s (exception class only above)",
                          job["id"])
            try:
                self.db.finalize_failed(job["id"], None)
                report.action = "aborted"
                report.detail["reason"] = "processing_error"
            except DoorError as e:
                log.error("abort finalize failed job=%s door=%s", job["id"], e.name)
                report.action = "error"
            return report

    # -- claimed-job processing --------------------------------------------
    def _process(self, job: dict, item_ids: list[str], report: RunReport) -> RunReport:
        rows = {r["id"].lower(): r for r in self.db.evidence_rows(item_ids)}
        if set(rows) != set(item_ids):
            raise RuntimeError("claimed items missing from evidence read")

        read_ids: list[str] = []
        failed: list[dict] = []
        contents: dict[str, str] = {}
        for iid in item_ids:
            row = rows[iid]
            if row["status"] == "deleted":
                # withdrawn after claim: still needs an explicit verdict —
                # the settle door records it as withdrawn
                read_ids.append(iid)
                continue
            try:
                contents[iid] = extract_content(row, self.storage)
                read_ids.append(iid)
            except ParseFailure as pf:
                failed.append({"item": iid, "code": pf.code})
        log.info(
            "read job=%s parsed=%d failed=%d", job["id"],
            len(contents), len(failed),
        )

        # Whole-batch input budget — checked BEFORE any epistemic commitment.
        total_chars = sum(len(c) for c in contents.values())
        if total_chars > BATCH_INPUT_CHAR_BUDGET:
            log.info("over budget job=%s chars=%d", job["id"], total_chars)
            self.db.finalize_failed(job["id"], FROZEN_CLIENT_COPY["batch_too_large"])
            report.action = "aborted"
            report.detail["reason"] = "batch_too_large"
            return report

        memory_entries: list[dict] = []
        reinforce_entries: list[dict] = []
        if contents:
            existing = self.db.comparison_memory(job["agent_id"])
            retracted = self.db.retracted_memory(job["agent_id"])
            # FAIL CLOSED on incomplete suppression context: synthesizing
            # while knowingly missing client retractions is forbidden. The
            # sentinel row (FETCH_LIMIT + 1) proves incompleteness; raising
            # here routes to the abort door via the standard error path.
            if len(retracted) > RETRACTED_FETCH_LIMIT:
                log.error(
                    "retracted context incomplete job=%s fetched=%d limit=%d",
                    job["id"], len(retracted), RETRACTED_FETCH_LIMIT,
                )
                raise RuntimeError("retracted_context_incomplete")
            log.info(
                "reconciliation job=%s existing_rows=%d retracted_rows=%d",
                job["id"], len(existing), len(retracted),
            )
            readable = [rows[i] for i in item_ids if i in contents]
            user_prompt = build_user_prompt(readable, contents, existing, retracted)

            raw = None
            error_hint = None
            for attempt in range(MODEL_MALFORMED_RETRIES + 1):
                prompt = user_prompt if error_hint is None else (
                    user_prompt
                    + f"\n\nYour previous output was invalid ({error_hint}). "
                    "Output ONLY the JSON object, exactly as specified."
                )
                raw = self.model.complete(SYSTEM_PROMPT, prompt)
                try:
                    memory_entries, reinforce_entries = validate_and_map(
                        raw, set(contents.keys()), existing, retracted
                    )
                    error_hint = None
                    break
                except ValidationError as ve:
                    error_hint = str(ve)
                    log.info(
                        "model output invalid job=%s attempt=%d reason=%s",
                        job["id"], attempt, error_hint,
                    )
            if error_hint is not None:
                # model could not produce a valid epistemic result: abort door
                self.db.finalize_failed(job["id"], None)
                report.action = "aborted"
                report.detail["reason"] = "model_output_invalid"
                return report

        # SETTLE — the only successful ending. Verdicts for EVERY claimed item.
        results = {
            "read": read_ids,
            "failed": failed,
            "memory": memory_entries,
            "reinforce": reinforce_entries,
        }
        settled = self.db.settle(job["id"], results, POLICY)
        log.info(
            "settled job=%s outcome=%s read=%s failed=%s withdrawn=%s "
            "inserted=%s tension=%s reinforced=%s discarded=%s dropped=%s "
            "grounded=%s agent_state=%s swept=%s",
            job["id"], settled.get("outcome"), settled.get("items_read"),
            settled.get("items_failed"), settled.get("items_withdrawn"),
            settled.get("memory_inserted"), settled.get("tension_rows"),
            settled.get("reinforced"), settled.get("statements_discarded"),
            settled.get("reinforcements_dropped"), settled.get("grounded_total"),
            settled.get("finalize", {}).get("agent_state"),
            settled.get("finalize", {}).get("swept_reading_items"),
        )
        report.action = "settled"
        report.outcome = settled.get("outcome")
        report.detail = {
            k: settled.get(k)
            for k in (
                "items_read", "items_failed", "items_withdrawn", "memory_inserted",
                "tension_rows", "reinforced", "statements_discarded",
                "reinforcements_dropped", "grounded_total",
            )
        }
        return report


# ---------------------------------------------------------------------------
# Entry point (production): env-configured, one job per invocation.
# ---------------------------------------------------------------------------

def main() -> int:
    import os

    logging.basicConfig(
        level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s"
    )
    # Public Actions logs carry ONLY the runner's contract lines: silence
    # third-party INFO chatter (the SDK's httpx request lines are harmless
    # but off-contract — observed in Fire #2 run 3).
    for noisy in ("httpx", "httpcore", "anthropic", "urllib3", "requests"):
        logging.getLogger(noisy).setLevel(logging.WARNING)
    base = os.environ["SUPABASE_URL"]
    key = os.environ["SUPABASE_SERVICE_KEY"]
    runner = Runner(
        db=RestDb(base, key),
        storage=SupabaseStorage(base, key),
        model=AnthropicModel(os.environ["ANTHROPIC_API_KEY"]),
    )
    report = runner.run_once()
    log.info(
        "done action=%s job=%s outcome=%s janitor=%d",
        report.action, report.job_id, report.outcome, report.janitor_finalized,
    )
    return 0 if report.action != "error" else 1


if __name__ == "__main__":
    raise SystemExit(main())
